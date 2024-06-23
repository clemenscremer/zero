import os
import pandas as pd
import mikeio
import matplotlib.pyplot as plt
import base64



# import paths from central config
from config import BASE_DIR, SIM_DATA_DIR, DOMAIN_DIR, INITIAL_DIR, SETUP_DIR, BOUNDARIES_DIR, RESULTS_DIR, FIGURE_DIR, TEMP_DIR, IMG_DIR
# import api client from central config
from config import client, DEPLOYMENT_NAME

# reporting
import datetime
from jinja2 import Template

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# --------------------------------------------------
# function definitions
# --------------------------------------------------
def list_model_files(return_folders=None):
    """ List all available model files in the specified subfolders of the 'sim_data' directory.
    Args:
    return_folders (list): List of folders whose content should be returned. Default is None, which returns content of all folders.

    Returns:
        dict: A dictionary containing the lists of files for each specified folder.
    """
    file_lists = {}

    # List files in the 'setup' subfolder
    if return_folders is None or 'setup' in return_folders:
        if os.path.exists(SETUP_DIR):
            setup_files = [file for file in os.listdir(SETUP_DIR) if file.endswith('.m21fm')]
            file_lists['setup'] = setup_files
        else:
            file_lists['setup'] = []

    # List files in the 'boundaries' subfolder
    if return_folders is None or 'boundaries' in return_folders:
        file_lists['boundaries'] = os.listdir(BOUNDARIES_DIR)

    # List files in the 'initial' subfolder
    if return_folders is None or 'initial' in return_folders:
        file_lists['initial'] = os.listdir(INITIAL_DIR)

    # List files in the 'domain' subfolder
    if return_folders is None or 'domain' in return_folders:
        file_lists['domain'] = os.listdir(DOMAIN_DIR)

    # List files in the 'results' subfolder
    if return_folders is None or 'result' in return_folders:
        file_lists['result'] = os.listdir(RESULTS_DIR)

    # List files in the 'figure' folder (contains results)
    if return_folders is None or 'figure' in return_folders:
        file_lists['figure'] = os.listdir(FIGURE_DIR)

    return file_lists


def get_pfs_parameters(file_path, start_time=True, time_step_interval=True, number_of_time_steps=True,
                       domain_file=True, initial_conditions_file=True, manning_number=True):
    """
    Retrieves specific parameters from a PFS numerical setup file.

    Args:
        file_path (str): Path to the PFS file, relative to the 'setup' directory.
        start_time (bool): Whether to retrieve the start time. Default is True.
        time_step_interval (bool): Whether to retrieve the time step interval. Default is True.
        number_of_time_steps (bool): Whether to retrieve the number of time steps. Default is True.
        domain_file (bool): Whether to retrieve the domain file name. Default is True.
        initial_conditions_file (bool): Whether to retrieve the initial conditions file name. Default is True.
        manning_number (bool): Whether to retrieve the Manning number. Default is True.

    Returns:
        dict: A dictionary containing the requested parameters.
    """
    from mikeio import read_pfs
    pfs = read_pfs(os.path.join(SETUP_DIR, file_path))
    parameters = {}

    if start_time:
        start_time_obj = pfs["FemEngineHD"]["TIME"]["start_time"]
        parameters["start_time"] = start_time_obj.strftime("%Y-%m-%d %H:%M:%S")  # Convert to string
    if time_step_interval:
        parameters["time_step_interval"] = pfs["FemEngineHD"]["TIME"]["time_step_interval"]
    if number_of_time_steps:
        parameters["number_of_time_steps"] = pfs["FemEngineHD"]["TIME"]["number_of_time_steps"]
    if domain_file:
        parameters["domain_file"] = pfs["FemEngineHD"]["DOMAIN"]["file_name"]
    if initial_conditions_file:
        parameters["initial_conditions_file"] = pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["INITIAL_CONDITIONS"]["file_name_2D"]
    if manning_number:
        parameters["manning_number"] = pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["BED_RESISTANCE"]["MANNING_NUMBER"]["constant_value"]

    return parameters


def modify_parameters(modified_parameters):
    """
    Modifies specific parameters in a PFS file, writes the modified file with a numbered suffix,
    and updates an Excel file with the simulation details.

    Args:
        modified_parameters (dict): Dictionary of parameters to modify, where the keys are the parameter names
                                  and the values are the new values.

    Returns:
        dict: Dictionary of parameters to modify, where the keys are the parameter names
                                  and the values are the new values.
    """
    return modified_parameters


def modify_pfs_parameters(file_path, modified_parameters):
    """
    Modifies specific parameters in a PFS file, writes the modified file with a numbered suffix,
    and updates an Excel file with the simulation details.

    Args:
        file_path (str): Path to the PFS file, relative to the 'setup' directory.
        modified_parameters (dict): Dictionary of parameters to modify, where the keys are the parameter names
                                  and the values are the new values.

    Returns:
        str: Path to the modified PFS file.
    """
    from mikeio import read_pfs

    pfs = read_pfs(os.path.join(SETUP_DIR, file_path))

    # Modify the specified parameters
    for parameter, value in modified_parameters.items():
        if parameter == "start_time":
            pfs["FemEngineHD"]["TIME"]["start_time"] = value
        elif parameter == "time_step_interval":
            pfs["FemEngineHD"]["TIME"]["time_step_interval"] = value
        elif parameter == "number_of_time_steps":
            pfs["FemEngineHD"]["TIME"]["number_of_time_steps"] = value
        elif parameter == "domain_file":
            pfs["FemEngineHD"]["DOMAIN"]["file_name"] = value
        elif parameter == "initial_conditions_file":
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["INITIAL_CONDITIONS"]["file_name_2D"] = value
        elif parameter == "initial_surface_elevation":
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["INITIAL_CONDITIONS"]["surface_elevation_constant"] = value
        elif parameter == "boundary_conditions_file":
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["BOUNDARY_CONDITIONS"]["CODE_2"]["file_name"] = value
        elif parameter == "manning_number":
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["BED_RESISTANCE"]["MANNING_NUMBER"]["constant_value"] = value
        else:
            raise ValueError(f"Parameter '{parameter}' is not a valid parameter in the PFS file.")


    # Get the list of existing setup files
    file_lists = list_model_files(return_folders=['setup'])
    setup_files = file_lists['setup']

    # Generate the modified file name with a numbered suffix
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    existing_numbers = [int(f.replace(base_name + "_", "").replace(".m21fm", "")) for f in setup_files if f.startswith(base_name + "_")]
    if existing_numbers:
        next_number = max(existing_numbers) + 1
    else:
        next_number = 1

    modified_file_path = os.path.join(SETUP_DIR, f"{base_name}_{next_number:03d}.m21fm")

    # Write the modified PFS file
    pfs.write(modified_file_path)

    # Update the Excel file with the simulation details
    excel_file_path = os.path.join(SIM_DATA_DIR, "simulations.xlsx")
    new_row = get_pfs_parameters(modified_file_path)
    new_row["setupfile_name"] = os.path.basename(modified_file_path)  # Add the file name to the new row
    if os.path.exists(excel_file_path):
        df = pd.read_excel(excel_file_path)
        df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
    else:
        df = pd.DataFrame([new_row])
    df.to_excel(excel_file_path, index=False)

    return modified_file_path


# --------------------------------------------------
def create_mesh_bathymetry(nx, ny, dx, dy, reference_depth, slope_degrees=0.0):
    """
    Create a mesh and bathymetry file.
    
    Parameters:
    nx (int): Number of cells in the x-direction.
    ny (int): Number of cells in the y-direction.
    dx (float): Cell size in the x-direction (in meters).
    dy (float): Cell size in the y-direction (in meters).
    reference_depth (float): Reference depth for the bathymetry (in meters).
    slope_degrees (float, optional): Slope of the bathymetry in degrees. Default is 0.0 (constant bathymetry).
    
    Returns:
    mikeio.Dataset: Bathymetry dataset.
    """
    import numpy as np
    import mikeio
    from mikeio import EUMType, EUMUnit


    if slope_degrees == 0.0:
        bathymetry = reference_depth * np.ones((ny, nx))
    else:
        slope = -np.tan(np.deg2rad(slope_degrees))
        x = np.arange(nx) * dx
        bathymetry = reference_depth - (x * slope)
        bathymetry = np.repeat(bathymetry[np.newaxis, :], ny, axis=0)
    
    dataset = mikeio.Dataset(
        [bathymetry],
        dims=("y", "x"),
        geometry=mikeio.Grid2D(
            nx=nx, ny=ny, dx=dx, dy=dy, x0=0.0, y0=0.0, projection="NON-UTM"
        ),
        items=[mikeio.ItemInfo("Bathymetry", EUMType.Bathymetry, EUMUnit.meter)],
    )

    bathymetry_filename = os.path.join(DOMAIN_DIR, f"mesh_bathy_{nx}x{ny}_dx{dx}_reference{reference_depth}_slope{slope_degrees}.dfs2")
    dataset.to_dfs(bathymetry_filename)

    return_str = f"Bathymetry dataset saved to: {bathymetry_filename}"
    return return_str


def plot_mesh_bathy(filename, plot_type='mesh_bathy'):
    """
    Plot the mesh and bathymetry or the initial surface elevation for a given filename in the 'domain' or 'initial' directory.

    Parameters:
    filename (str): The name of the file containing the mesh and bathymetry or initial surface elevation data.
    plot_type (str, optional): Whether to plot the 'mesh_bathy' or 'initial' conditions. Default is 'mesh_bathy'.

    Returns:
    matplotlib.figure.Figure: A matplotlib figure object representing the plot of the mesh and bathymetry or the initial surface elevation.
    """
    import matplotlib.pyplot as plt
    if plot_type == 'mesh_bathy':
        file_path = os.path.join(DOMAIN_DIR, filename)
    elif plot_type == 'initial':
        file_path = os.path.join(INITIAL_DIR, filename)
    else:
        raise ValueError("Invalid plot_type. Must be 'mesh_bathy' or 'initial'.")

    ds = mikeio.read(file_path)

    # Calculate the scaling factor to maintain aspect ratio
    scy = ds.geometry.y.max() / ds.geometry.x.max()

    fig, ax = plt.subplots(figsize=(6, 6 * scy))
    if plot_type == 'mesh_bathy':
        ds["Bathymetry"].plot(ax=ax, label="z [m]")
        ax.set_title(f"Mesh and Bathymetry: {filename}")
    elif plot_type == 'initial':
        ds["Surface elevation"].plot(ax=ax, label="z [m]")
        ax.set_title(f"Initial Surface Elevation: {filename}")

    return fig


def create_surface_elevation(nx, ny, dx, dy, wave_height=0.0, wave_width=0.0, wave_position='left'): 
    """ Create a surface elevation dataset.
    Parameters:
    nx (int): Number of cells in the x-direction.
    ny (int): Number of cells in the y-direction.
    dx (float): Cell size in the x-direction (in meters).
    dy (float): Cell size in the y-direction (in meters).
    wave_height (float, optional): Height of the sinusoidal surface elevation wave (in meters). Default is 0.0 (no wave).
    wave_width (float, optional): Width of the sinusoidal surface elevation wave (in meters). Default is 0.0 (no wave).
    wave_position (str, optional): Position of the sinusoidal surface elevation wave, either 'left', 'right', 'center', or 'both'. Default is 'left'.

    Returns:
    mikeio.Dataset: Surface elevation dataset.
    """
    import numpy as np
    from mikeio import EUMType, EUMUnit

    surface_elevation = np.zeros((ny, nx))

    if wave_height > 0.0 and wave_width > 0.0:
        x = np.arange(nx) * dx
        n_wave_points = int(wave_width / dx)
        x_wave = np.linspace(0, 1, n_wave_points)
        y_wave = np.sin(np.pi * x_wave) * wave_height

        if wave_position == 'left':
            surface_elevation[:, :n_wave_points] = y_wave
        elif wave_position == 'right':
            surface_elevation[:, -n_wave_points:] = y_wave
        elif wave_position == 'center':
            surface_elevation[:, nx//2-n_wave_points//2:nx//2+n_wave_points//2] = y_wave
        elif wave_position == 'both':
            surface_elevation[:, :n_wave_points] = y_wave
            surface_elevation[:, -n_wave_points:] = y_wave

    dataset = mikeio.Dataset(
        [surface_elevation],
        dims=("y", "x"),
        geometry=mikeio.Grid2D(
            nx=nx, ny=ny, dx=dx, dy=dy, x0=0.0, y0=0.0, projection="NON-UTM"
        ),
        items=[mikeio.ItemInfo("Surface elevation", EUMType.Surface_Elevation, EUMUnit.meter)],
    )
    
    surface_elevation_filename = os.path.join(INITIAL_DIR, 
    f"surface_elevation_{nx}x{ny}_dx{dx}wave{wave_height}x{wave_width}{wave_position}.dfs2") 
    dataset.to_dfs(surface_elevation_filename) 
    return_str = f"Initial surface elevation dataset saved to: {surface_elevation_filename}"

    return return_str

def save_setup(file_path, modified_parameters):
    """
    Modifies a PFS file with the given parameters and saves it with a numbered suffix.

    Args:
        file_path (str): Path to the original PFS file, relative to the 'setup' directory.
        modified_parameters (dict): Dictionary of parameters to modify, where the keys are the parameter names
                                    and the values are the new values.

    Returns:
        str: Path to the newly saved PFS file.
    """
    from mikeio import read_pfs

    # Read the original PFS file
    pfs = read_pfs(os.path.join(SETUP_DIR, file_path))

    # Modify the specified parameters
    for parameter, value in modified_parameters.items():
        if parameter == "start_time":
            pfs["FemEngineHD"]["TIME"]["start_time"] = value
        elif parameter == "time_step_interval":
            pfs["FemEngineHD"]["TIME"]["time_step_interval"] = value
        elif parameter == "number_of_time_steps":
            pfs["FemEngineHD"]["TIME"]["number_of_time_steps"] = value
        elif parameter == "domain_file":
            #pfs["FemEngineHD"]["DOMAIN"]["file_name"] = f"|{value}|"
            pfs["FemEngineHD"]["DOMAIN"]["file_name"] = value

        elif parameter == "initial_conditions_file":
            #pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["INITIAL_CONDITIONS"]["file_name_2D"] = f"|{value}|"
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["INITIAL_CONDITIONS"]["file_name_2D"] = value

        elif parameter == "manning_number":
            pfs["FemEngineHD"]["HYDRODYNAMIC_MODULE"]["BED_RESISTANCE"]["MANNING_NUMBER"]["constant_value"] = value
        else:
            raise ValueError(f"Parameter '{parameter}' is not a valid parameter in the PFS file.")

    # Get the list of existing setup files
    file_lists = list_model_files(return_folders=['setup'])
    setup_files = file_lists['setup']

    # Generate the modified file name with a numbered suffix
    base_name = os.path.splitext(os.path.basename(file_path))[0]
    existing_numbers = [int(f.replace(base_name + "_", "").replace(".m21fm", "")) for f in setup_files if f.startswith(base_name + "_")]
    if existing_numbers:
        next_number = max(existing_numbers) + 1
    else:
        next_number = 1

    modified_file_path = os.path.join(SETUP_DIR, f"{base_name}_{next_number:03d}.m21fm")

    # Write the modified PFS file
    pfs.write(modified_file_path)

    return modified_file_path

def simulate(simfile_name):
    """Simulate MIKE Model using the specified simfile."""
    import mikesimulation
    sim_path = os.path.join(SETUP_DIR, simfile_name)
    mikesimulation.execute_simulation(sim_path, verbose=True)

    try:
        # copy simulation results from SETUP_DIR subfolder simfile_name, " - Result Files" to RESULTS_DIR whilst renaming to simfile_name.dfsu
        #/teamspace/studios/this_studio/mk-assistant/sim_data/setup/sim_.m21fm/ - Result Files/area.dfsu'
        import shutil
        results_in_path = os.path.join(SETUP_DIR, f"{simfile_name} - Result Files", "area.dfsu")
        results_out_path = os.path.join(RESULTS_DIR, simfile_name.replace(".m21fm", ".dfsu"))
        shutil.copy(results_in_path, results_out_path)
        # remove old results folder
        shutil.rmtree(os.path.join(SETUP_DIR, f"{simfile_name} - Result Files"))
        return f"Simulation ran successfully, results copied to {results_out_path}"

    except Exception as e:
        return f"Error copying simulation results: {e}. Likely simulation error"
    

def plot_results(result_file, n_times=3):
    """
    plot results of simulation
    arguments:
    result_file: string, dfsu result file to create plot from
    n_times: number of time steps to plot
    """
    import mikeio
    import matplotlib.pyplot as plt
    import numpy as np

    ds = mikeio.read(os.path.join(RESULTS_DIR, result_file))
    total_timesteps = len(ds.time)
    # catch if n_times is larger than total timesteps
    if n_times > total_timesteps:
        n_times = total_timesteps
    
    fig, ax = plt.subplots(n_times, 1, figsize=(7, 2.3*n_times))
    for i, time in enumerate(np.linspace(0, total_timesteps-1, n_times, dtype=int)):
        if n_times == 1:
            ds["Surface elevation"].plot(ax=ax)
        else:
            ds["Surface elevation"].isel(time=time).plot(ax=ax[i])
    
    fig.suptitle(f"Simulation data from {result_file}")
    output_fn = result_file.replace(".dfsu", ".png")
    fig.savefig(os.path.join(FIGURE_DIR, result_file.replace(".dfsu", ".png")))
    return fig #f"Plots saved to {FIGURE_DIR} as {output_fn}"



# Evaluation
def encode_image(image_path):
    """Encodes an image to a base64 string."""
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

def analyze_images(image_filenames, added_context):
    """
    Analyzes one or more images using GPT-4.

    Parameters:
    image_filenames (list): A list of image filenames to be analyzed.
    added_context (str): Additional context to provide to the LLM, e.g. parameters of simulations.

    Returns:
    str: The response message from GPT-4.
    """
    
    # Ensure image_filenames is a list even if a single filename is provided
    if isinstance(image_filenames, str):
        image_filenames = [image_filenames]

    # Full paths to images
    image_paths = [os.path.join(FIGURE_DIR, filename) for filename in image_filenames]

    # Encode images and prepare them for the prompt
    base64_images = [encode_image(image_path) for image_path in image_paths]

        
    sys_prompt = """
    You are an engineering assistant specializing in water-related projects, focusing on numerical simulations and equipped to analyze images of 
    2d simulation results. 
    You will be asked to analyze images of water-related simulations and provide a brief, concise description of the results.
    Describe features such as minima and maxima, trends, and patterns are of interest.
    Expect to contrast the image with other images e.g. from other timesteps of the simulation or other simulations.

    Exemplary descriptions:
    * "Water levels range from 0 to 5 meters. A wave is located at x=100m and skewed to the right."
    * "The wave propagation speeds are similar in both, but the wave dissipates faster in the second simulation."

    (optional) additional context:
    """

    if added_context is not None:
        sys_prompt += added_context

    messages = [
        {"role": "system", "content": sys_prompt},
    ]

    for base64_image in base64_images:
        messages.append({
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {
                    "url": f"data:image/png;base64,{base64_image}"}
                }
            ]
        })

    # Call AI
    response = client.chat.completions.create(
        model=DEPLOYMENT_NAME,
        messages=messages,
    )
    response_message = response.choices[0].message
    # Create a matplotlib figure to display the analyzed images
    fig, axes = plt.subplots(1, len(image_filenames), figsize=(12, 4))
    # Ensure axes is an array-like object, even if there's only one image
    if not hasattr(axes, '__iter__'):
        axes = [axes]

    for i, filename in enumerate(image_filenames):
        axes[i].imshow(plt.imread(os.path.join(FIGURE_DIR, filename)))
        axes[i].set_title(filename)
        axes[i].axis('off')
    plt.tight_layout()
    fig.savefig("test_analysis.png")
    return {
        "response_message": response_message,
        "figure": fig
    }

def format_markdown(doc, text):
    # Split the text into lines
    lines = text.split('\n')
    for line in lines:
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('• ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            if '**' in line:
                parts = line[2:].split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:  # Odd parts are bold
                        p.add_run(part).bold = True
                    else:
                        p.add_run(part)
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd parts are bold
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        elif '`' in line:
            p = doc.add_paragraph()
            parts = line.split('`')
            for i, part in enumerate(parts):
                if i % 2 == 1:  # Odd parts are code
                    p.add_run(part).font.name = 'Courier New'
                else:
                    p.add_run(part)
        elif ':' in line:
            key, value = line.split(':', 1)
            p = doc.add_paragraph()
            p.add_run(key.strip() + ":").bold = True
            p.add_run(value.strip())
        else:
            doc.add_paragraph(line)

def generate_report(messages, image_filenames):
    # System prompt for the report generation
    system_prompt = """
    You are an AI assistant tasked with creating a comprehensive engineering report based on a chat interaction about numerical simulations. Your report should include:

    1. Introduction: Provide a short outline and motivation for the study. If not explicitly stated, infer the purpose based on the tasks performed and questions asked.

    2. Methods: Summarize the simulation setup, including domains, initial conditions, simulation durations, and time-stepping details.

    3. Results: Summarize the results of the simulation(s) or comparisons between different simulations. Reference the provided figures in your discussion using [FIGURE1], [FIGURE2], etc.

    Format your response using the following markdown-like syntax:
    - Use '# ' for main sections (Introduction, Methods, Results, Conclusion)
    - Use '## ' for subsections
    - Use '### ' for sub-subsections
    - Use '#### ' for sub-sub-sections
    - Use '• ' or '- ' for bullet points
    - Use '1. ' for numbered lists
    - Use '**bold text**' for emphasis
    - Use '`code`' for file names or technical terms
    - Use 'Key: Value' format for parameter listings

    Ensure consistent formatting throughout the report. For simulation parameters, use bullet points with bold keys, like this:
    • **Parameter Name:** Value

    For file names, use the `code` format.
    """

    # Prepare the chat history for the LLM
    chat_history = [msg for msg in messages if msg['role'] != 'system']

    # Prepare image information for the LLM
    image_info = [f"Figure {i+1}: Results from simulation {filename}" for i, filename in enumerate(image_filenames)]
    image_context = "Available figures:\n" + "\n".join(image_info)

    # Call the LLM to generate the report content
    response = client.chat.completions.create(
        model=DEPLOYMENT_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Generate a report based on this chat history: {chat_history}\n\n{image_context}"}
        ],
        temperature=0.5,
    )

    report_content = response.choices[0].message.content

    # Create a new Document
    doc = Document()

    # Update existing styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    # Create a new style for the title
    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.base_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(18)
    title_style.font.bold = True

    # Add title
    doc.add_paragraph('Numerical Simulation Report', style='CustomTitle')

    # Process and add the LLM-generated content
    sections = report_content.split('\n\n')
    for section in sections:
        format_markdown(doc, section)

    # Add images
    for i, img_filename in enumerate(image_filenames, 1):
        img_path = os.path.join(FIGURE_DIR, img_filename)
        doc.add_picture(img_path, width=Inches(6))
        caption = doc.add_paragraph(f'Figure {i}: Results from simulation {img_filename}')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Replace [FIGURE] placeholders in the content
        for paragraph in doc.paragraphs:
            if f'[FIGURE{i}]' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'[FIGURE{i}]', f'(See Figure {i})')

    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add DHI logo to the header
    header = section.header
    header_para = header.paragraphs[0]
    logo_run = header_para.add_run()
    logo_path = os.path.join(IMG_DIR, "DHI_Logo_Pos_Black_noMargin.png")
    logo_run.add_picture(logo_path, width=Inches(1))

    # Save the document
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_filename = f'simulation_report_{timestamp}.docx'
    docx_path = os.path.join(TEMP_DIR, docx_filename)
    doc.save(docx_path)

    return {
        "message": "Report generated successfully",
        "file_path": docx_path
    }

# --------------------------------------------------
# List of descriptions for all functions
# --------------------------------------------------
function_descriptions = []


function_descriptions.append(
    {
        "name": "list_model_files",
        "description": "List all files in the 'boundaries', 'initial', 'domain', and 'setup', 'result' and 'figure' subfolders of the simulation data directory.",
        "parameters": {
            "type": "object",
            "properties": {
                "return_folders": {
                    "type": "array",
                    "description": "List of folders whose content should be returned. Can be used to keep return small, e.g. the user is solely interested in setup. Default is None, which returns content of all folders.",
                    "items": {
                        "type": "string"
                    }
                }
            },
            "required": []
        },
        "returns": {
            "type": "object",
            "description": "A dictionary containing the lists of files for each subfolder specified in return_folders.",
            "properties": {
                "boundaries": {
                    "type": "array",
                    "description": "List of files in the 'boundaries' subfolder."
                },
                "initial": {
                    "type": "array",
                    "description": "List of files in the 'initial' subfolder."
                },
                "domain": {
                    "type": "array",
                    "description": "List of files in the 'domain' subfolder."
                },
                "setup": {
                    "type": "array",
                    "description": "List of *.m21fm files in the 'setup' subfolder."
                },
                "result": {
                    "type": "array",
                    "description": "List of files in the 'result' subfolder."
                },
                "figure": {
                    "type": "array",
                    "description": "List of files in the 'figure' subfolder."
                }
            }
        }
    },
    )


function_descriptions.append(
    {
        "name": "get_pfs_parameters",
        "description": "Retrieves specific parameters from a PFS numerical setup file.",
        "parameters": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the PFS file, relative to the 'setup' directory."
                },
                "start_time": {
                    "type": "boolean",
                    "description": "Whether to retrieve the start time. Default is True."
                },
                "time_step_interval": {
                    "type": "boolean",
                    "description": "Whether to retrieve the time step interval. Default is True."
                },
                "number_of_time_steps": {
                    "type": "boolean",
                    "description": "Whether to retrieve the number of time steps. Default is True."
                },
                "domain_file": {
                    "type": "boolean",
                    "description": "Whether to retrieve the domain file name. Default is True. Contains Bathymetry and mesh"
                },
                "initial_conditions_file": {
                    "type": "boolean",
                    "description": "Whether to retrieve the initial conditions file name. Default is True."
                },
                "manning_number": {
                    "type": "boolean",
                    "description": "Whether to retrieve the Manning number. Default is True."
                }
            },
            "required": [
                "file_path"
            ]
        },
        "returns": {
            "type": "object",
            "description": "A dictionary containing the requested parameters.",
            "properties": {
                "start_time": {
                    "type": "string",
                    "description": "The start time of the simulation."
                },
                "time_step_interval": {
                    "type": "number",
                    "description": "The time step interval in seconds."
                },
                "number_of_time_steps": {
                    "type": "integer",
                    "description": "The number of time steps in the simulation."
                },
                "domain_file": {
                    "type": "string",
                    "description": "The file name of the domain file (mesh and bathymetry)."
                },
                "initial_conditions_file": {
                    "type": "string",
                    "description": "The file name of the initial conditions file."
                },
                "manning_number": {
                    "type": "number",
                    "description": "The constant Manning number."
                }
            }
        }
    }
)


function_descriptions.append(
    {
        "name": "modify_parameters",
        "description": "Prepares a dictionary of parameters to modify in a PFS file based on user input.",
        "parameters": {
            "type": "object",
            "properties": {
                "modified_parameters": {
                    "type": "object",
                    "description": "Dictionary of parameters to modify, where the keys are the parameter names and the values are the new values.",
                    "properties": {
                        "start_time": {
                            "type": "string",
                            "description": "The new start time of the simulation."
                        },
                        "time_step_interval": {
                            "type": "number",
                            "description": "The new time step interval in seconds."
                        },
                        "number_of_time_steps": {
                            "type": "integer",
                            "description": "The new number of time steps in the simulation."
                        },
                        "domain_file": {
                            "type": "string",
                            "description": "The new file name of the domain file (mesh and bathymetry)."
                        },
                        "initial_conditions_file": {
                            "type": "string",
                            "description": "The new file name of the initial conditions file."
                        },
                        "manning_number": {
                            "type": "number",
                            "description": "The new constant Manning number."
                        }
                    }
                }
            },
            "required": [
                "modified_parameters"
            ]
        },
        "returns": {
            "type": "object",
            "description": "The dictionary of parameters to modify."
        }
    })


    
function_descriptions.append(
    {
        "name": "create_mesh_bathymetry",
        "description": "Create a 2D mesh and bathymetry dataset for a numerical simulation and save it to a file.",
        "parameters": {
            "type": "object",
            "properties": {
                "nx": {
                    "type": "integer",
                    "description": "Number of cells in the x-direction."
                },
                "ny": {
                    "type": "integer",
                    "description": "Number of cells in the y-direction."
                },
                "dx": {
                    "type": "number",
                    "description": "Cell size in the x-direction (in meters)."
                },
                "dy": {
                    "type": "number",
                    "description": "Cell size in the y-direction (in meters)."
                },
                "reference_depth": {
                    "type": "number",
                    "description": "Reference depth for the bathymetry (in meters)."
                },
                "slope_degrees": {
                    "type": "number",
                    "description": "Slope of the bathymetry in degrees. Default is 0.0 (constant bathymetry).",
                    "default": 0.0
                }
            },
            "required": [
                "nx",
                "ny",
                "dx",
                "dy",
                "reference_depth"
            ]
        },
        "returns": {
            "type": "string",
            "description": "The file path of the saved bathymetry dataset."
        }
}
)

function_descriptions.append(
{
    "name": "plot_mesh_bathy",
    "description": "Plot the mesh and bathymetry or the initial surface elevation for a given filename in the 'domain' or 'initial' directory.",
    "parameters": {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "The name of the file containing the mesh and bathymetry or initial surface elevation data."
            },
            "plot_type": {
                "type": "string",
                "description": "Whether to plot the 'mesh_bathy' or 'initial' conditions. Default is 'mesh_bathy'.",
                "default": "mesh_bathy",
                "enum": [
                    "mesh_bathy",
                    "initial"
                ]
            }
        },
        "required": [
            "filename"
        ]
    },
    "returns": {
        "type": "object",
        "description": "A matplotlib.figure.Figure object representing the plot of the mesh and bathymetry or the initial surface elevation."
    }
}
)

function_descriptions.append(
 
 {
    "name": "create_surface_elevation",
    "description": "Create a 2D surface elevation dataset for the initial conditions of a numerical simulation and save it to a file.",
    "parameters": {
        "type": "object",
        "properties": {
            "nx": {
                "type": "integer",
                "description": "Number of cells in the x-direction."
            },
            "ny": {
                "type": "integer",
                "description": "Number of cells in the y-direction."
            },
            "dx": {
                "type": "number",
                "description": "Cell size in the x-direction (in meters)."
            },
            "dy": {
                "type": "number",
                "description": "Cell size in the y-direction (in meters)."
            },
            "wave_height": {
                "type": "number",
                "description": "Height of the sinusoidal surface elevation wave (in meters). Default is 0.0 (no wave).",
                "default": 0.0
            },
            "wave_width": {
                "type": "number",
                "description": "Width of the sinusoidal surface elevation wave (in meters). Default is 0.0 (no wave).",
                "default": 0.0
            },
            "wave_position": {
                "type": "string",
                "description": "Position of the sinusoidal surface elevation wave, either 'left', 'right', 'center', or 'both'. Default is 'left'.",
                "default": "left",
                "enum": [
                    "left",
                    "right",
                    "center",
                    "both"
                ]
            }
        },
        "required": [
            "nx",
            "ny",
            "dx",
            "dy"
        ]
    },
    "returns": {
        "type": "string",
        "description": "The file path of the saved surface elevation dataset."
    }
}
)

function_descriptions.append(
    {
        "name": "save_setup",
        "description": "Modifies a PFS file with the given parameters and saves it with a numbered suffix.",
        "parameters": {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Path to the original PFS file, relative to the 'setup' directory."
                },
                "modified_parameters": {
                    "type": "object",
                    "description": "Dictionary of parameters to modify, where the keys are the parameter names and the values are the new values.",
                    "properties": {
                        "start_time": {
                            "type": "string",
                            "description": "The new start time of the simulation."
                        },
                        "time_step_interval": {
                            "type": "number",
                            "description": "The new time step interval in seconds."
                        },
                        "number_of_time_steps": {
                            "type": "integer",
                            "description": "The new number of time steps in the simulation."
                        },
                        "domain_file": {
                            "type": "string",
                            "description": "The new file name of the domain file (mesh and bathymetry)."
                        },
                        "initial_conditions_file": {
                            "type": "string",
                            "description": "The new file name of the initial conditions file."
                        },
                        "initial_surface_elevation": {
                            "type": "number",
                            "description": "The new constant surface elevation in meters."
                        },
                        "boundary_conditions_file": {
                            "type": "string",
                            "description": "The new file name of the boundary conditions file."
                        },
                        "manning_number": {
                            "type": "number",
                            "description": "The new constant Manning number."
                        }
                    }
                }
            },
            "required": [
                "file_path",
                "modified_parameters"
            ]
        },
        "returns": {
            "type": "string",
            "description": "The path to the newly saved PFS file."
        }
    })
function_descriptions.append(
    {
        "name": "simulate",
        "description": "Simulate MIKE Model using the specified simfile. E.g. if the user asks to execute simulation or run simulation",
        "parameters": {
            "type": "object",
            "properties": {
                "simfile_name": {
                    "type": "string",
                    "description": "The name of the simulation file *.m21fm."
                }
            },
            "required": [
                "simfile_name"
            ]
        },
        "returns": {
            "type": "string",
            "description": "A string indicating the status of the simulation."
        }
    }
    )


function_descriptions.append(
    {
        "name": "plot_results",
        "description": "Plot results of a simulation from a DFSU result file",
        "parameters": {
            "type": "object",
            "properties": {
                "result_file": {
                    "type": "string",
                    "description": "The name of the DFSU result file to create the plot from"
                },
                "n_times": {
                    "type": "integer",
                    "description": "The number of time steps to plot",
                    "default": 3
                }
            },
            "required": [
                "result_file"
            ]
        },
        "returns": {
            "type": "object",
            "description": "A matplotlib.figure.Figure object representing the plot of the results with the specified number of time steps, saved to the 'figures' subfolder."
        }
    }
)

function_descriptions.append(
    {
        "name": "analyze_images",
        "description": "Analyzes one or more images using GPT-4o and provides a brief, concise description of the results, as well as a matplotlib figure showing the analyzed images side-by-side.",
        "parameters": {
            "type": "object",
            "properties": {
                "image_filenames": {
                    "type": "array",
                    "description": "A list of image filenames to be analyzed. There will be an image for each simulation that can comprise multiple timesteps. E.g. named sim__001.png",
                    "items": {
                        "type": "string"
                    }
                },
                "added_context": {
                    "type": "string",
                    "description": "Additional context to provide to the LLM for analysis, e.g. parameters of simulations which can partly be derived from filenames if not known from convarsation."
                }
            },
            "required": [
                "image_filenames"
            ]
        },
        "returns": {
            "type": "object",
            "description": "An object containing the response message from GPT-4o and a matplotlib figure object displaying the analyzed images side-by-side.",
            "properties": {
                "response_message": {
                    "type": "string",
                    "description": "The response message from GPT-4o containing the analysis of the provided images."
                },
                "figure": {
                    "type": "object",
                    "description": "A matplotlib figure object displaying the analyzed images side-by-side."
                }
            }
        }
    }
)


function_descriptions.append(
    {
        "name": "generate_report",
        "description": "Generate a comprehensive report based on the chat history and simulation results.",
        "parameters": {
            "type": "object",
            "properties": {
                "messages": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "role": {"type": "string"},
                            "content": {"type": "string"}
                        }
                    },
                    "description": "List of message dictionaries from the session state."
                },
                "image_filenames": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of image filenames to include in the report."
                }
            },
            "required": ["messages", "image_filenames"]
        },
        "returns": {
            "type": "string",
            "description": "Path to the generated HTML report."
        }
    }
)

# --------------------------------------------------
# define available functions. 
# NOTE: Only those will be visible to LLM, hence can 
# quickly be switched on/off here 
# --------------------------------------------------
available_functions = {
    # os
    "list_model_files": list_model_files,
    # setup
    "get_pfs_parameters": get_pfs_parameters,

    "modify_parameters": modify_parameters,


    "create_mesh_bathymetry": create_mesh_bathymetry,

    "create_surface_elevation": create_surface_elevation,

    # plotting
    "plot_mesh_bathy": plot_mesh_bathy,

    "plot_results": plot_results,

    "save_setup": save_setup,

    # simulation
    "simulate": simulate,

    # evaluate
    "analyze_images": analyze_images,

    #
    "generate_report": generate_report,

}